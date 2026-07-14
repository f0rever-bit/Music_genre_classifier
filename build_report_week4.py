from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

SRC = "Project_Report_Week3.docx"
OUT = "Project_Report_Week4.docx"

doc = Document(SRC)

# ---- Clear existing body (paragraphs + tables) ----
body = doc.element.body
for child in list(body):
    body.remove(child)

def title_line(text, size_pt, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER):
    p = doc.add_paragraph()
    p.alignment = align
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size_pt)
    return p

def h1(text):
    return doc.add_paragraph(text, style="Heading 1")

def h2(text):
    return doc.add_paragraph(text, style="Heading 2")

def body_para(text):
    return doc.add_paragraph(text, style="Normal")

def toc_line(text, style="toc 1"):
    return doc.add_paragraph(text, style=style)

def bullet(bold_prefix, rest):
    p = doc.add_paragraph(style="List Bullet")
    r1 = p.add_run(bold_prefix)
    r1.bold = True
    p.add_run(rest)
    return p

def spacer():
    doc.add_paragraph()

# ======================== TITLE PAGE ========================
title_line("Національний Університет Біоресурсів і Природокористування", 24)
title_line("Дисципліна: Навчальна технологічна практика", 24)
spacer()
title_line("Назва проєкту: Audio-Based Music Recommender", 24)
title_line("(Music Genre Classifier)", 16, bold=False)
spacer()
title_line("Тиждень 4: Фінальний звіт про виконання", 18)
spacer(); spacer()
title_line("Автор: студент", 15, align=WD_ALIGN_PARAGRAPH.RIGHT)
title_line("Групи ІПЗ-24008б", 15, align=WD_ALIGN_PARAGRAPH.RIGHT)
title_line("Федченко Андрій", 15, align=WD_ALIGN_PARAGRAPH.RIGHT)

# ======================== ЗМІСТ ========================
spacer()
h1("Зміст")
spacer()
toc_line("Зміст\t2")
toc_line("1. Виконані завдання (Тиждень 4)\t3", "toc 1")
toc_line("1.1 Завершені етапи розробки\t3", "toc 2")
toc_line("1.2 Тестування та якість\t4", "toc 2")
toc_line("1.3 Інфраструктура та інтеграції\t4", "toc 2")
toc_line("2. Аналіз та проблеми\t5", "toc 1")
toc_line("2.1 Технічні виклики\t5", "toc 2")
toc_line("2.2 Знайдені рішення (цикл вдосконалення)\t5", "toc 2")
toc_line("3. Оновлення архітектури та бази даних\t6", "toc 1")
toc_line("4. Фінальний стан проєкту\t8", "toc 1")
toc_line("4.1 Досягнуті результати\t8", "toc 2")
toc_line("4.2 Можливі шляхи розвитку (Future Work)\t8", "toc 2")
toc_line("5. Додаткова інформація\t10", "toc 1")
toc_line("5.1 Технологічний стек\t10", "toc 2")
toc_line("5.2 Статистика проєкту\t10", "toc 2")

# ======================== 1. ВИКОНАНІ ЗАВДАННЯ ========================
spacer()
h1("1. Виконані завдання (Тиждень 4)")
body_para(
    "Протягом четвертого тижня роботи над проєктом Audio-Based Music Recommender "
    "було реалізовано найважливіші фічі з беклогу, які забезпечують повноцінну "
    "функціональність продукту. Особливу увагу приділено інтеграції зі Spotify, "
    "папкам для користувачів, A/B тестуванню та жанровій обізнаності ML-моделі."
)

h2("1.1 Завершені етапи розробки")
bullet("Крок 1 — Папки та Slug-роутинг: ",
       "Додано модель Folder (папки) для групування треків користувачів. "
       "Впроваджено людинозрозумілі URL-адреси (slug-роутинг) на базі виконавця та назви треку "
       "замість числових ID. Реалізовано повний цикл управління папками на бекенді та фронтенді.")
bullet("Крок 2 — Інтеграція зі Spotify: ",
       "Створено гібридну модель даних (Spotify + In-Memory). Реалізовано глобальний плеєр "
       "(Spotify Web Playback SDK) для Premium-користувачів із fallback-варіантом на 30-секундні "
       "прев'ю для Free-користувачів. Додано підтримку обкладинок (cover art) зі Spotify.")
bullet("Крок 3 — Жанрово-усвідомлені рекомендації: ",
       "Вектор ознак для K-Means розширено: додано інформацію про жанр (one-hot encoding на 25+ жанрів). "
       "Додано механізм автоматичного перенавчання моделі при значному розширенні бази треків.")
bullet("Крок 4 — A/B тестування та Admin-панель: ",
       "Реалізовано статистичний розрахунок значущості результатів A/B тестування (z-test). "
       "Створено дашборд адміністратора для перегляду статистики, кількості користувачів та зручної "
       "можливості зробити алгоритм-переможець алгоритмом за замовчуванням.")
bullet("Крок 5 — Покращення AI-тегування: ",
       "Замінено нестабільний API MusicBrainz на iTunes Search API. Тепер автоматичне визначення "
       "жанрів і метаданих через Gemini + iTunes працює стабільно і дозволило успішно перетегувати 700+ треків.")
bullet("Крок 6 — База даних та безпека: ",
       "Оновлено bcrypt до версії 4.0.1 із захистом від довгих паролів. Додано нативні PostgreSQL "
       "тригери для поля updated_at. Додано нові композитні індекси для оптимізації аналітики.")

h2("1.2 Тестування та якість")
body_para(
    "Додано багато нових тестів, що підвищило надійність та стабільність додатку. "
    "Покриття коду розширено на нові сервіси та ендпоінти."
)
bullet("Загальна кількість тестів: ", "190 (значне зростання після 3-го тижня).")
bullet("Покриття (Coverage): ", "Покриття бекенду сягнуло 82%.")
bullet("Нові тести: ", "Кешування (Redis), Storage-абстракції, Audio utils, Mock-сервіси, міграції БД, "
       "A/B статистика та Spotify API.")

h2("1.3 Інфраструктура та інтеграції")
bullet("Гібридна модель: ", "Зберігання та аналіз локальних файлів у RAM з подальшим видаленням (для безкоштовного хостингу), "
       "а також підключення зовнішнього каталогу (Spotify).")
bullet("OAuth: ", "Вдосконалено OAuth авторизацію Spotify із можливістю примусової зміни акаунта (show_dialog=true).")

# ======================== 2. АНАЛІЗ ТА ПРОБЛЕМИ ========================
spacer()
h1("2. Аналіз та проблеми")
body_para(
    "Четвертий тиждень був зосереджений на інтеграції зовнішніх сервісів (Spotify, iTunes, Gemini) "
    "та покращенні алгоритмів рекомендацій. Нижче описано основні виклики."
)

h2("2.1 Технічні виклики")
bullet("Обмеження Spotify Web Playback SDK: ",
       "Плеєр вимагав Premium-підписку. Для користувачів без Premium відтворення блокувалось. "
       "Рішення: переписано логіку GlobalPlayer для автоматичного переходу на відтворення 30-секундного "
       "прев'ю (HTML5 <audio>), якщо SDK повертає помилку 'not_connected' / 'premium_required'.")
bullet("Проблеми з метаданими (MusicBrainz): ",
       "Безкоштовний API MusicBrainz часто повертав порожні жанри або 'Jazz/Rock' для всіх треків. "
       "Рішення: замінено основне джерело жанрів на iTunes Search API, що підвищило точність класифікації.")
bullet("Локальні файли та місце на сервері: ",
       "Зберігання всіх MP3 файлів на диску є дорогим. "
       "Рішення: опція DELETE_LOCAL_AFTER_ANALYZE, що дозволяє робити аналіз у RAM і видаляти "
       "фізичний файл одразу після отримання аудіоознак.")
bullet("Точність K-Means: ",
       "Кластеризація раніше ігнорувала стиль музики, опираючись лише на тембр і темп. "
       "Рішення: жанр перетворено у вектор і додано до загального feature-вектора з відповідною вагою (GENRE_WEIGHT = 3.0).")

h2("2.2 Знайдені рішення (цикл вдосконалення)")
bullet("", "Додано підтримку 'Слагів' (human-readable URLs) замість ID в ендпоінтах.")
bullet("", "Вирішено проблеми з CORS для локальної розробки React-додатку.")
bullet("", "Додано Redis кеш (TTL 5 min) для швидкої віддачі рекомендацій.")
bullet("", "Налаштовано автоматичне перенавчання K-Means моделі (auto_retrain_if_needed).")

# ======================== 3. ОНОВЛЕННЯ АРХІТЕКТУРИ ========================
spacer()
h1("3. Оновлення архітектури та бази даних")
body_para("Під час 4-го тижня було внесено ряд критичних змін до схеми бази даних через Alembic міграції:")
bullet("Міграція 007: ", "Додано таблицю ABConfig для збереження алгоритму рекомендацій за замовчуванням.")
bullet("Міграція 008: ", "Додано PostgreSQL тригери для поля updated_at (ON UPDATE).")
bullet("Міграція 009: ", "Додано композитні B-tree індекси для AlgorithmEvent.")
bullet("Міграція 010 (Hybrid Source): ", "Поле file_path стало необов'язковим. Додано поля "
       "source, external_id, preview_url, stream_url. Змінено унікальні індекси (partial unique indexes).")
bullet("Міграція 013: ", "Додано поле cover_url для зберігання обкладинок треків.")
bullet("Міграція 014: ", "Додано таблицю Folders. Додано поле folder_id та slug до таблиці Music.")

# ======================== 4. ФІНАЛЬНИЙ СТАН ПРОЄКТУ ========================
spacer()
h1("4. Фінальний стан проєкту")
h2("4.1 Досягнуті результати")
body_para(
    "Проєкт 'Audio-Based Music Recommender' є повністю функціональним MVP. "
    "Виконано 100% запланованих задач. Система підтримує завантаження аудіо, автоматичний аналіз "
    "аудіоознак (Librosa), AI-тегування (Gemini + iTunes), кластеризацію (Scikit-Learn), A/B тестування "
    "та відтворення музики через інтеграцію зі Spotify SDK."
)

h2("4.2 Можливі шляхи розвитку (Future Work)")
bullet("Cloud Deploy: ", "Розгортання на платформі типу Render або AWS.")
bullet("Експорт у Spotify: ", "Автоматичне створення плейлиста у Spotify із знайдених рекомендацій.")
bullet("Соціальні функції: ", "Шерінг папок (публічні URL), профілі користувачів.")
bullet("Багатомовність (i18n): ", "Повноцінне перемикання мов фронтенду без перезавантаження сторінки.")

# ======================== 5. ДОДАТКОВА ІНФОРМАЦІЯ ========================
spacer()
h1("5. Додаткова інформація")

h2("5.1 Технологічний стек")
bullet("Backend: ", "Python 3.10+, FastAPI, SQLAlchemy 2.0, PostgreSQL / SQLite, Pydantic v2, "
       "librosa, scikit-learn, joblib, slowapi, python-dotenv.")
bullet("Frontend: ", "React 18, Vite, JavaScript, Axios, react-plotly.js.")
bullet("Інтеграції: ", "Spotify Web API / Web Playback SDK, iTunes Search API, Google Gemini API (gemini-2.5-flash).")
bullet("DevOps: ", "Docker, Docker Compose, Nginx, Alembic, GitHub Actions.")
bullet("Тестування: ", "pytest (190 тестів).")

h2("5.2 Статистика проєкту")
bullet("Кількість тестів: ", "190 (82% покриття).")
bullet("Кількість моделей БД: ", "7 (User, Music, AudioFeatures, Recommendation, AlgorithmEvent, ABConfig, Folder).")
bullet("Міграцій Alembic: ", "14.")
bullet("Прогрес: ", "100% — MVP готовий до захисту.")

doc.save(OUT)
print("Saved", OUT)
