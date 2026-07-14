from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

SRC = "Project_Report_Week4.docx"
OUT = "Project_Report_Week5.docx"

doc = Document(SRC)

# ---- Clear existing body (paragraphs + tables) ----
body = doc.element.body
for child in list(body):
    body.remove(child)

def title_line(text, size_pt, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER):
    p = doc.add_paragraph()
    p.alignment = align
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size_pt)
    return p

def h1(text):
    return doc.add_paragraph(text, style="Heading 1")

def h2(text):
    return doc.add_paragraph(text, style="Heading 2")

def body_para(text):
    return doc.add_paragraph(text, style="Normal")

def toc_line(text, style="toc 1"):
    return doc.add_paragraph(text, style=style)

def bullet(bold_prefix, rest):
    p = doc.add_paragraph(style="List Bullet")
    r1 = p.add_run(bold_prefix)
    r1.bold = True
    p.add_run(rest)
    return p

def spacer():
    doc.add_paragraph()

# ======================== TITLE PAGE ========================
title_line("Національний Університет Біоресурсів і Природокористування", 24)
title_line("Дисципліна: Навчальна технологічна практика", 24)
spacer()
title_line("Назва проєкту: Audio-Based Music Recommender", 24)
title_line("(Music Genre Classifier)", 16, bold=False)
spacer()
title_line("Тиждень 5: Експорт у Spotify та інтеграція API", 18)
spacer(); spacer()
title_line("Автор: студент", 15, align=WD_ALIGN_PARAGRAPH.RIGHT)
title_line("Групи ІПЗ-24008б", 15, align=WD_ALIGN_PARAGRAPH.RIGHT)
title_line("Федченко Андрій", 15, align=WD_ALIGN_PARAGRAPH.RIGHT)

# ======================== ЗМІСТ ========================
spacer()
h1("Зміст")
spacer()
toc_line("Зміст\t2")
toc_line("1. Виконані завдання (Тиждень 5)\t3", "toc 1")
toc_line("1.1 Завершені етапи розробки\t3", "toc 2")
toc_line("1.2 Інтеграція зі Spotify API\t4", "toc 2")
toc_line("2. Аналіз та проблеми\t5", "toc 1")
toc_line("2.1 Технічні виклики\t5", "toc 2")
toc_line("3. Фінальний стан проєкту\t7", "toc 1")
toc_line("4. Додаткова інформація\t8", "toc 1")
toc_line("5. Додаткова інформація\t10", "toc 1")
toc_line("5.1 Технологічний стек\t10", "toc 2")
toc_line("5.2 Статистика проєкту\t10", "toc 2")

# ======================== 1. ВИКОНАНІ ЗАВДАННЯ ========================
spacer()
h1("1. Виконані завдання (Тиждень 5)")
body_para(
    "Протягом п'ятого тижня було реалізовано довгоочікувану фічу з розділу Future Work: "
    "експорт знайдених рекомендацій у приватний плейліст Spotify. Також розширено "
    "ліміти рекомендацій для зручнішого використання масового експорту."
)

h2("1.1 Завершені етапи розробки")
bullet("Експорт у Spotify: ",
       "Додано кнопку 'Export to Spotify' на сторінці рекомендацій. Створено ендпоінт "
       "для автоматичного створення плейліста і наповнення його треками через Spotify API.")
bullet("Масове вивантаження: ",
       "Збільшено ліміт запиту до рекомендацій з 50 до 100 треків, щоб користувач міг "
       "генерувати та зберігати великі плейлісти за один клік.")
bullet("Обробка скоупів OAuth: ",
       "Розширено список прав (scopes) у Spotify-інтеграції на playlist-modify-public та "
       "playlist-modify-private.")

h2("1.2 Інтеграція зі Spotify API")
body_para(
    "Для створення плейлістів застосовано новітні стандарти Spotify Web API. Зокрема, "
    "було здійснено міграцію з застарілих ендпоінтів."
)

# ======================== 2. АНАЛІЗ ТА ПРОБЛЕМИ ========================
spacer()
h1("2. Аналіз та проблеми")
body_para("Під час впровадження експорту виникла серія непередбачуваних помилок API.")

h2("2.1 Технічні виклики")
bullet("Deprecation ендпоінту створення плейлістів: ",
       "Під час тестування сервер повертав 403 Forbidden. З'ясувалося, що з лютого 2026 року "
       "Spotify зробив застарілим ендпоінт POST /v1/users/{user_id}/playlists. "
       "Рішення: переписано логіку під новий ендпоінт POST /v1/me/playlists.")
bullet("Deprecation ендпоінту додавання треків: ",
       "Також був видалений ендпоінт POST /v1/playlists/{playlist_id}/tracks. "
       "Рішення: код бекенду оновлено для використання нового шляху POST /v1/playlists/{playlist_id}/items.")

# ======================== 3. ФІНАЛЬНИЙ СТАН ПРОЄКТУ ========================
spacer()
h1("3. Фінальний стан проєкту")
body_para(
    "Проєкт тепер повністю інтегровано зі Spotify не лише як плеєр чи джерело інформації, "
    "але й як місце збереження згенерованих ШІ-плейлістів."
)
bullet("Експорт: ", "Повноцінно працює експорт пачок до 100 треків.")

# ======================== 5. ДОДАТКОВА ІНФОРМАЦІЯ ========================
spacer()
h1("4. Додаткова інформація")

h2("4.1 Технологічний стек")
bullet("Backend: ", "Python 3.10+, FastAPI, SQLAlchemy 2.0, PostgreSQL / SQLite, Pydantic v2, "
       "librosa, scikit-learn, joblib, slowapi, python-dotenv.")
bullet("Frontend: ", "React 18, Vite, JavaScript, Axios, react-plotly.js.")
bullet("Інтеграції: ", "Spotify Web API / Web Playback SDK, iTunes Search API, Google Gemini API (gemini-2.5-flash).")
bullet("DevOps: ", "Docker, Docker Compose, Nginx, Alembic, GitHub Actions.")
bullet("Тестування: ", "pytest (190 тестів).")

h2("4.2 Статистика проєкту")
bullet("Кількість тестів: ", "190 (82% покриття).")
bullet("Кількість моделей БД: ", "7 (User, Music, AudioFeatures, Recommendation, AlgorithmEvent, ABConfig, Folder).")
bullet("Міграцій Alembic: ", "14.")
bullet("Прогрес: ", "100% — MVP готовий до захисту.")

doc.save(OUT)
print("Saved", OUT)
